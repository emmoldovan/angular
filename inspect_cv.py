from docx import Document
from docx.shared import Pt, Inches, Emu
import json

doc = Document(r"c:\Users\Emanuel\Personals\interviuri\Angular\Resume Emanuel MOLDOVAN - 2.12.2025.docx")

# Section/page margins
for i, section in enumerate(doc.sections):
    print(f"\n=== SECTION {i} ===")
    print(f"  Top margin: {section.top_margin} ({section.top_margin / 914400:.2f} inches)")
    print(f"  Bottom margin: {section.bottom_margin} ({section.bottom_margin / 914400:.2f} inches)")
    print(f"  Left margin: {section.left_margin} ({section.left_margin / 914400:.2f} inches)")
    print(f"  Right margin: {section.right_margin} ({section.right_margin / 914400:.2f} inches)")
    print(f"  Page width: {section.page_width} ({section.page_width / 914400:.2f} inches)")
    print(f"  Page height: {section.page_height} ({section.page_height / 914400:.2f} inches)")

# Headers/footers
for i, section in enumerate(doc.sections):
    header = section.header
    if header and header.paragraphs:
        print(f"\n=== HEADER SECTION {i} ===")
        for p in header.paragraphs:
            print(f"  Header text: '{p.text}'")
            for run in p.runs:
                print(f"    Run: '{run.text}' bold={run.bold} size={run.font.size} color={run.font.color.rgb if run.font.color and run.font.color.rgb else 'None'}")

print("\n\n=== STYLES USED ===")
styles_used = set()
for p in doc.paragraphs:
    styles_used.add(p.style.name)
for s in sorted(styles_used):
    print(f"  {s}")

print("\n\n=== ALL PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    text_preview = p.text[:80] if p.text else "(empty)"
    align = p.alignment
    spacing_before = p.paragraph_format.space_before
    spacing_after = p.paragraph_format.space_after
    line_spacing = p.paragraph_format.line_spacing
    indent_left = p.paragraph_format.left_indent

    print(f"\n--- Para {i}: style='{p.style.name}' align={align} ---")
    print(f"  Text: '{text_preview}'")
    print(f"  Space before={spacing_before}, after={spacing_after}, line_spacing={line_spacing}")
    print(f"  Left indent={indent_left}")

    for j, run in enumerate(p.runs):
        run_text = run.text[:60] if run.text else "(empty)"
        print(f"  Run {j}: '{run_text}'")
        print(f"    bold={run.bold}, italic={run.italic}, underline={run.underline}")
        print(f"    font.name='{run.font.name}', size={run.font.size}")
        if run.font.color and run.font.color.rgb:
            print(f"    color={run.font.color.rgb}")

# Check for tables
print(f"\n\n=== TABLES: {len(doc.tables)} ===")
for i, table in enumerate(doc.tables):
    print(f"\nTable {i}: {len(table.rows)} rows x {len(table.columns)} cols")
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            text = cell.text[:60]
            print(f"  Cell [{ri},{ci}]: '{text}'")
            for p in cell.paragraphs:
                print(f"    Para style='{p.style.name}' align={p.alignment}")
                for run in p.runs:
                    print(f"      Run: '{run.text[:40]}' bold={run.bold} italic={run.italic} size={run.font.size} color={run.font.color.rgb if run.font.color and run.font.color.rgb else None}")
