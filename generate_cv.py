"""
Clone the original CV and make targeted modifications to match JD requirements.
Preserves all original formatting, fonts, spacing, and styles.
"""
import copy
from docx import Document
from docx.shared import Pt, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

INPUT = r"c:\Users\Emanuel\Personals\interviuri\Angular\Resume Emanuel MOLDOVAN - 2.12.2025.docx"
OUTPUT = r"c:\Users\Emanuel\Personals\interviuri\Angular\Resume Emanuel MOLDOVAN - UPDATED2.docx"

doc = Document(INPUT)

# ============================================================
# HELPERS
# ============================================================

def get_para_text(p):
    return p.text.strip()

def find_para_index(doc, substring):
    """Find paragraph index by substring match."""
    for i, p in enumerate(doc.paragraphs):
        if substring in p.text:
            return i
    return -1

def clone_paragraph_after(doc, source_para, new_text, index_after):
    """
    Clone a paragraph's XML (preserving style, formatting, numbering),
    replace its text with new_text, and insert it after the given index.
    """
    # Deep copy the source paragraph XML
    new_p_xml = copy.deepcopy(source_para._p)

    # Clear all runs and replace with single run containing new_text
    # But preserve the formatting from the first run
    runs = new_p_xml.findall(qn('w:r'))
    if runs:
        # Keep first run's formatting, remove the rest
        first_run = runs[0]
        for r in runs[1:]:
            new_p_xml.remove(r)
        # Update text in first run
        t_elements = first_run.findall(qn('w:t'))
        if t_elements:
            t_elements[0].text = new_text
            t_elements[0].set(qn('xml:space'), 'preserve')
        else:
            t = OxmlElement('w:t')
            t.text = new_text
            t.set(qn('xml:space'), 'preserve')
            first_run.append(t)
        # Remove bold from the run if present (bullets shouldn't be bold)
        rPr = first_run.find(qn('w:rPr'))
        if rPr is not None:
            b = rPr.find(qn('w:b'))
            if b is not None:
                rPr.remove(b)

    # Insert after the target paragraph
    target_p = doc.paragraphs[index_after]._p
    target_p.addnext(new_p_xml)

    return new_p_xml

def replace_para_text(para, old_substr, new_substr):
    """Replace text within a paragraph, preserving run formatting."""
    for run in para.runs:
        if old_substr in run.text:
            run.text = run.text.replace(old_substr, new_substr)
            return True
    # If substring spans multiple runs, do full text replacement
    full_text = para.text
    if old_substr in full_text:
        new_full = full_text.replace(old_substr, new_substr)
        # Set all text in first run, clear rest
        runs = para.runs
        if runs:
            runs[0].text = new_full
            for r in runs[1:]:
                r.text = ''
            return True
    return False

def insert_bullet_after(doc, ref_index, text, template_para):
    """Insert a new bullet point after ref_index, cloned from template_para."""
    return clone_paragraph_after(doc, template_para, text, ref_index)


# ============================================================
# 1. SKILLS BAR - Add Module Federation, Azure DevOps, Bicep
# ============================================================
# Para 3 is the skills bar
skills_para = doc.paragraphs[3]
# Add new run at the end with same formatting as existing runs
last_run = skills_para.runs[-1]
new_run = skills_para.add_run('\n| Module Federation / Webpack 5 | Azure DevOps | Azure Bicep')
new_run.bold = True
new_run.font.name = 'Cambria'
# Size is inherited (None) in original

print("[OK] Skills bar updated")

# ============================================================
# 2. SUMMARY - Add autonomous collaboration, clean code, agile
# ============================================================
# Para 5 is the summary. It has a bold first part and regular second part.
# We need to modify the regular text runs.
summary_para = doc.paragraphs[5]
# The summary spans multiple runs. Let's reconstruct the non-bold part.
# Keep the bold run as-is, replace the rest.
bold_text = "Reliable and adaptable Front-End Engineer with 10+ years of experience "
new_body = (
    "building and optimizing scalable, high-performance web and mobile applications. "
    "Proven ability to collaborate autonomously with cross-functional teams (backend, product, design) "
    "and rapidly master new technologies. Deep expertise in modern web "
    "technologies (JavaScript, TypeScript, React, Vue.js, Lit, Node.js). "
    "Committed to delivering clean, well-structured code at a sustained and predictable pace "
    "within Agile/Scrum workflows. Passionate about joining an innovative software company "
    "that values staying at the forefront of web trends and continuous learning."
)

# Set first non-bold run to have all the new text, clear the rest
runs = summary_para.runs
found_bold_end = False
first_normal_set = False
for run in runs:
    if run.bold:
        continue
    if not first_normal_set:
        run.text = new_body
        first_normal_set = True
    else:
        run.text = ''

print("[OK] Summary updated")

# ============================================================
# 3. ADOBE - "Autonomously collaborated" + Agile bullet
# ============================================================
# Para 15 = "Collaborated closely with product, design..."
adobe_collab_idx = find_para_index(doc, "Collaborated closely with product, design")
if adobe_collab_idx >= 0:
    replace_para_text(doc.paragraphs[adobe_collab_idx],
                      "Collaborated closely with product, design, and backend teams via Jira and Figma to ship high-impact features on time",
                      "Autonomously collaborated with product, design, and backend teams via Jira and Figma to ship high-impact features on time")
    print(f"[OK] Adobe collaboration updated (para {adobe_collab_idx})")

    # Add Agile bullet after Adobe's last bullet
    # Find the empty para after Adobe bullets (para 16)
    agile_text = "Actively participated in daily standups, sprint plannings, and retrospectives, ensuring alignment and predictable delivery across sprints."
    insert_bullet_after(doc, adobe_collab_idx, agile_text, doc.paragraphs[adobe_collab_idx])
    print("[OK] Adobe Agile bullet added")

# ============================================================
# 4. HAIILO - Add Agile bullet
# ============================================================
haiilo_bullet_idx = find_para_index(doc, "Developed and maintained a comprehensive library of Stencil UI")
if haiilo_bullet_idx >= 0:
    agile_text = "Actively participated in daily standups, sprint plannings, and retrospectives."
    insert_bullet_after(doc, haiilo_bullet_idx, agile_text, doc.paragraphs[haiilo_bullet_idx])
    print(f"[OK] Haiilo Agile bullet added (after para {haiilo_bullet_idx})")

# ============================================================
# 5. AROBS [SCOT] - Add CI/CD, Bicep, Agile bullets
# ============================================================
scot_last_bullet_idx = find_para_index(doc, "Collaborated closely with backend and management teams, actively participating")
if scot_last_bullet_idx >= 0:
    template = doc.paragraphs[scot_last_bullet_idx]

    # Insert in reverse order (each goes after the same index, so last inserted = first in doc)
    agile_text = "Actively participated in daily standups, sprint plannings, and retrospectives."
    insert_bullet_after(doc, scot_last_bullet_idx, agile_text, template)

    bicep_text = "Provisioned and managed Azure cloud infrastructure using Bicep (Infrastructure as Code) for consistent and repeatable deployments."
    insert_bullet_after(doc, scot_last_bullet_idx, bicep_text, template)

    cicd_text = "Configured and managed CI/CD pipelines on Azure DevOps, automating build, test, and deployment workflows for consistent and reliable releases."
    insert_bullet_after(doc, scot_last_bullet_idx, cicd_text, template)

    print(f"[OK] Arobs SCOT bullets added (CI/CD, Bicep, Agile)")

# ============================================================
# 6. AROBS [CLIVE] - Micro-frontend + BFF + autonomous + Agile
# ============================================================
# Refresh paragraph indices after insertions above
clive_micro_idx = find_para_index(doc, "Implement a scalable solution by using microservices architecture")
if clive_micro_idx >= 0:
    replace_para_text(doc.paragraphs[clive_micro_idx],
                      "Implement a scalable solution by using microservices architecture and Angular libraries benefits",
                      "Implemented a scalable solution using micro-frontend architecture with Module Federation (Webpack 5) and Angular libraries")
    print(f"[OK] CLIVE microservices -> micro-frontend (para {clive_micro_idx})")

# Add BFF bullet after the micro-frontend bullet
clive_micro_idx = find_para_index(doc, "micro-frontend architecture with Module Federation")
if clive_micro_idx >= 0:
    bff_text = "Implemented and maintained a BFF (Back-end for Front-end) layer in C# / .NET Core, aggregating data from multiple microservices to serve the Angular frontend."
    insert_bullet_after(doc, clive_micro_idx, bff_text, doc.paragraphs[clive_micro_idx])
    print("[OK] CLIVE BFF bullet added")

# Add autonomous collaboration bullet after "Worked closely with product team"
clive_product_idx = find_para_index(doc, "Worked closely with the product team to digitize all processes")
if clive_product_idx >= 0:
    autonomous_text = "Acted as the autonomous frontend point of contact for client\u2019s internal teams (backend, product, other micro-frontend squads)."
    insert_bullet_after(doc, clive_product_idx, autonomous_text, doc.paragraphs[clive_product_idx])
    print("[OK] CLIVE autonomous bullet added")

# Add Agile bullet after mentoring bullet
clive_mentor_idx = find_para_index(doc, "Effectively mentored 4 junior developers")
if clive_mentor_idx >= 0:
    agile_text = "Actively participated in daily standups, sprint plannings, and retrospectives."
    insert_bullet_after(doc, clive_mentor_idx, agile_text, doc.paragraphs[clive_mentor_idx])
    print("[OK] CLIVE Agile bullet added")

# ============================================================
# 7. IVANTI - Add Agile bullet
# ============================================================
ivanti_last_idx = find_para_index(doc, "Create new applications that integrate into the ecosystem")
if ivanti_last_idx >= 0:
    agile_text = "Actively participated in daily standups, sprint plannings, and retrospectives."
    insert_bullet_after(doc, ivanti_last_idx, agile_text, doc.paragraphs[ivanti_last_idx])
    print(f"[OK] Ivanti Agile bullet added")

# ============================================================
# 8. MOVALIO - Vue Composition API + Agile
# ============================================================
movalio_vue_idx = find_para_index(doc, "big vending company in Romania")
if movalio_vue_idx >= 0:
    replace_para_text(doc.paragraphs[movalio_vue_idx],
                      "using Vue and Vuex",
                      "using Vue (Composition API) and Vuex")
    print(f"[OK] Movalio Vue -> Composition API")

movalio_last_idx = find_para_index(doc, "Contributed to an in-house web app that manage")
if movalio_last_idx >= 0:
    agile_text = "Actively participated in daily standups, sprint plannings, and retrospectives."
    insert_bullet_after(doc, movalio_last_idx, agile_text, doc.paragraphs[movalio_last_idx])
    print("[OK] Movalio Agile bullet added")

# ============================================================
# 9. ROBOYO - Add Agile bullet
# ============================================================
roboyo_idx = find_para_index(doc, "Created a web app for an UiPath collaborator")
if roboyo_idx >= 0:
    agile_text = "Actively participated in daily standups, sprint plannings, and retrospectives."
    insert_bullet_after(doc, roboyo_idx, agile_text, doc.paragraphs[roboyo_idx])
    print("[OK] Roboyo Agile bullet added")

# ============================================================
# 10. CLINICDR - Add Agile bullet
# ============================================================
clinicdr_last_idx = find_para_index(doc, "Participated in every step of the product development process")
if clinicdr_last_idx >= 0:
    agile_text = "Actively participated in daily standups, sprint plannings, and retrospectives."
    insert_bullet_after(doc, clinicdr_last_idx, agile_text, doc.paragraphs[clinicdr_last_idx])
    print("[OK] ClinicDr Agile bullet added")

# ============================================================
# 11. PROFESSIONAL SKILLS - Add BFF, CI/CD, Agile, clean code
# ============================================================
# Modify existing "Good back-end experience" bullet
backend_skill_idx = find_para_index(doc, "Good back-end experience")
if backend_skill_idx >= 0:
    p = doc.paragraphs[backend_skill_idx]
    # Clear all runs and set new text in first run
    runs = p.runs
    full_new = "Good back-end experience (.NET Core, Node.js) including BFF pattern implementation"
    runs[0].text = full_new
    for r in runs[1:]:
        r.text = ''
    print("[OK] Professional Skills - backend updated")

# Add new skills after "Outstanding organizational..."
outstanding_idx = find_para_index(doc, "Outstanding organizational")
if outstanding_idx >= 0:
    template = doc.paragraphs[outstanding_idx]

    # Insert in reverse order
    clean_text = "Strong focus on clean code, code reviews, and predictable delivery cadence"
    insert_bullet_after(doc, outstanding_idx, clean_text, template)

    agile_skill = "Active practitioner of Agile/Scrum ceremonies (daily standups, sprint planning, retrospectives)"
    insert_bullet_after(doc, outstanding_idx, agile_skill, template)

    cicd_skill = "Experience with CI/CD pipelines (Azure DevOps) and Infrastructure as Code (Azure Bicep)"
    insert_bullet_after(doc, outstanding_idx, cicd_skill, template)

    print("[OK] Professional Skills - 3 new bullets added")

# ============================================================
# SAVE
# ============================================================
doc.save(OUTPUT)
print(f"\nCV saved to: {OUTPUT}")
